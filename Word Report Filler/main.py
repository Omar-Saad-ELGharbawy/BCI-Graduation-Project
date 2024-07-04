import json
from docx import Document
from docx2pdf import convert
import os
import pyrebase

firebase_config = {
    "apiKey": "AIzaSyBaf2PZ9lRnkpM952pBDlfGCxxEjcvU4Bk",
    "authDomain": "fes-controller.firebaseapp.com",
    "databaseURL": "https://fes-controller-default-rtdb.firebaseio.com",
    "projectId": "fes-controller",
    "storageBucket": "fes-controller.appspot.com",
    "messagingSenderId": "341771619246",
    "appId": "1:341771619246:web:29d7cc61c889b3228a163f",
}


def replace_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{key}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{key}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{key}}}", str(value))


def create_filled_docx(template_path, output_path, data):
    doc = Document(template_path)
    replace_placeholders(doc, data)
    doc.save(output_path)


def upload_to_firebase(file_path, firebase_config):
    firebase = pyrebase.initialize_app(firebase_config)
    storage = firebase.storage()
    file_name = os.path.basename(file_path)
    storage.child(file_name).put(file_path)
    print(f"File uploaded to {storage.child(file_name).get_url(None)}")


def main():
        template_path = "ReportTemplate.docx"
        output_docx_path = "FilledReport.docx"
        json_path = "data.json"    

        with open(json_path, "r") as json_file:
            data = json.load(json_file)

        output_pdf_path = f"{data["Name"]}_{data["Session_Date"]}.pdf"

        # Create the docx
        create_filled_docx(template_path, output_docx_path, data)

        # Convert the filled Word document to PDF
        convert(output_docx_path, output_pdf_path)

        # Upload the PDF to Firebase Storage
        upload_to_firebase(output_pdf_path, firebase_config)

        # Delete the DOCX file after conversion
        if os.path.exists(output_docx_path):
            os.remove(output_docx_path)
            os.remove(output_pdf_path)

main()
