import json
from docx import Document
from docx2pdf import convert
import os
import pyrebase

firebase_config = {
    "apiKey": "AIzaSyBaf2PZ9lRnkpM952pBDlfGCxxEjcvU4Bk",
    "authDomain": "fes-controller.firebaseapp.com",
    "databaseURL": "https://fes-controller-default-rtdb.firebaseio.com",
    "projectId": "fes-controller",
    "storageBucket": "fes-controller.appspot.com",
    "messagingSenderId": "341771619246",
    "appId": "1:341771619246:web:29d7cc61c889b3228a163f",
}


class Report:
    def __init__(self, template_path, json_path):
        self.template_path = template_path
        self.json_path = json_path
        with open(json_path, "r") as json_file:
            data = json.load(json_file)

        self.data = data

    def replace_placeholders(self, doc):
        for paragraph in doc.paragraphs:
            for key, value in self.data.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in self.data.items():
                        if f"{{{key}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{key}}}", str(value))

    def create_filled_docx(self, output_docx_path):
        doc = Document(self.template_path)
        self.replace_placeholders(doc)
        doc.save(output_docx_path)

    def convert_to_pdf(self, output_docx_path, output_pdf_path):
        convert(output_docx_path, output_pdf_path)

    def upload_to_firebase_storage(self, file_path):
        firebase = pyrebase.initialize_app(firebase_config)
        storage = firebase.storage()
        file_name = os.path.basename(file_path)
        storage.child(file_name).put(file_path)
        print(f"File uploaded to {storage.child(file_name).get_url(None)}")

    def upload_to_firebase_database(self):
        firebase = pyrebase.initialize_app(firebase_config)
        db = firebase.database()
        db.child("reports").push(self.to_dict())

    def generate_report(self):
        output_docx_path = "Word Report Filler\FilledReport.docx"
        output_pdf_path = (
            f"Word Report Filler\{self.data['Name']}_{self.data['Session_Date']}.pdf"
        )

        self.create_filled_docx(output_docx_path)
        self.convert_to_pdf(output_docx_path, output_pdf_path)
        self.upload_to_firebase_storage(output_pdf_path)

        # Delete the DOCX file after conversion
        if os.path.exists(output_docx_path):
            os.remove(output_docx_path)
            os.remove(output_pdf_path)


if __name__ == "__main__":
    template_path = "Word Report Filler\ReportTemplate.docx"
    json_path = "Word Report Filler\data.json"

    report = Report(template_path, json_path)
    report.generate_report()
