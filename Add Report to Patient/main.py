import os
import pyrebase
from docx import Document
from docx2pdf import convert

config = {
    "apiKey": "AIzaSyBaf2PZ9lRnkpM952pBDlfGCxxEjcvU4Bk",
    "authDomain": "fes-controller.firebaseapp.com",
    "databaseURL": "https://fes-controller-default-rtdb.firebaseio.com",
    "projectId": "fes-controller",
    "storageBucket": "fes-controller.appspot.com",
    "messagingSenderId": "341771619246",
    "appId": "1:341771619246:web:29d7cc61c889b3228a163f",
}

# Initialize Firebase
firebase = pyrebase.initialize_app(config)
db = firebase.database()
storage = firebase.storage()

TEMPLATE_PATH = "Add Report to Patient/ReportTemplate.docx"
output_docx_path = "Add Report to Patient/FilledReport.docx"


# Define the patient data class
class Patient:
    def __init__(
        self,
        name,
        age,
        gender,
        phone_number,
        mobile_number,
        address,
        national_id,
        height,
        weight,
        dominant_hand,
        affected_hand,
        disease,
        object_shape,
    ):
        self.name = name
        self.age = age
        self.gender = gender
        self.phone_number = phone_number
        self.mobile_number = mobile_number
        self.address = address
        self.national_id = national_id
        self.height = height
        self.weight = weight
        self.dominant_hand = dominant_hand
        self.affected_hand = affected_hand
        self.disease = disease
        self.object_shape = object_shape
        self.reports = []

    def to_dict(self):
        return {
            "Name": self.name,
            "Age": self.age,
            "Gender": self.gender,
            "PhoneNum": self.phone_number,
            "MobileNum": self.mobile_number,
            "Address": self.address,
            "NationalId": self.national_id,
            "Height": self.height,
            "Weight": self.weight,
            "DominantHand": self.dominant_hand,
            "AffectedHand": self.affected_hand,
            "Disease": self.disease,
            "ObjectShape": self.object_shape,
            "reports": [report.to_dict() for report in self.reports],
        }

    def upload_patient_to_firebase(self):
        db.child("patients").child(self.name).set(self.to_dict())

    def add_report(self, report):
        self.reports.append(report)  # Append the Report instance


class Report:
    def __init__(
        self,
        Name="",
        Patient_ID="",
        Date_of_Birth="",
        Gender="",
        Hand_Dominance="",
        Diagnosis="",
        Referring_Clinician="",
        Rehabilitation_Start_Date="",
        Muscle_Weakness="",
        Muscle_Paralysis="",
        Muscle_Spasticity="",
        Loss_of_Coordination="",
        Tremors="",
        Reduced_Range_of_Motion="",
        Sensory_Loss="",
        Pain="",
        Hand_function_impairment="",
        Muscle_Testing_Score="",
        Number_of_Sessions="",
        Discharge_Date="",
        Medical_History="",
        Waveform="",
        Pulse_Frequency="",
        Pulse_Intensity="",
        Pulse_Width="",
        Movement_Angle_cup="",
        Accuracy_cup="",
        Feedback_cup="",
        is_adjustment_Required_cup="",
        Additional_Notes_cup="",
        Movement_Angle_pen="",
        Accuracy_pen="",
        Feedback_pen="",
        is_adjustment_Required_pen="",
        Additional_Notes_pen="",
        Movement_Angle_key="",
        Accuracy_key="",
        Feedback_key="",
        is_adjustment_Required_key="",
        Additional_Notes_key="",
        Movement_Angle_ball="",
        Accuracy_ball="",
        Feedback_ball="",
        is_adjustment_Required_ball="",
        Additional_Notes_ball="",
        Movement_Angle_bottle="",
        Accuracy_bottle="",
        Feedback_bottle="",
        is_adjustment_Required_bottle="",
        Additional_Notes_bottle="",
        Session_Date="",
        Stimulation_Protocol="",
        Target_Movements="",
        Session_Duration="",
        Performance_Feedback="",
        Clinician_Notes="",
        Improvements_Observed="",
        Challenges_Noted="",
        Side_effects="",
        Next_Steps_in_Therapy="",
        Continue_Current_Protocol="",
        Adjust_Stimulation_Settings="",
        New_Settings="",
        Increase_Session_Frequency="",
        Introduce_New_Exercises="",
        Details="",
        experience_satisfying="",
        comfortable_therapy="",
        session_painful="",
        hand_improved="",
        Additional_Comments="",
        Quality_of_life_Improved="",
        Long_Term_Goals_Achieved="",
        Remaining_Challenges="",
        Continue_Home_Exercises="",
        Further_Clinical_Sessions_Needed="",
        Additional_Therapies_Suggested="",
    ):
        self.data = {
            "Name": Name,
            "Patient_ID": Patient_ID,
            "Date_of_Birth": Date_of_Birth,
            "Gender": Gender,
            "Hand_Dominance": Hand_Dominance,
            "Diagnosis": Diagnosis,
            "Referring_Clinician": Referring_Clinician,
            "Rehabilitation_Start_Date": Rehabilitation_Start_Date,
            "Muscle_Weakness": Muscle_Weakness,
            "Muscle_Paralysis": Muscle_Paralysis,
            "Muscle_Spasticity": Muscle_Spasticity,
            "Loss_of_Coordination": Loss_of_Coordination,
            "Tremors": Tremors,
            "Reduced_Range_of_Motion": Reduced_Range_of_Motion,
            "Sensory_Loss": Sensory_Loss,
            "Pain": Pain,
            "Hand_function_impairment": Hand_function_impairment,
            "Muscle_Testing_Score": Muscle_Testing_Score,
            "Number_of_Sessions": Number_of_Sessions,
            "Discharge_Date": Discharge_Date,
            "Medical_History": Medical_History,
            "Waveform": Waveform,
            "Pulse_Frequency": Pulse_Frequency,
            "Pulse_Intensity": Pulse_Intensity,
            "Pulse_Width": Pulse_Width,
            "Movement_Angle_cup": Movement_Angle_cup,
            "Accuracy_cup": Accuracy_cup,
            "Feedback_cup": Feedback_cup,
            "is_adjustment_Required_cup": is_adjustment_Required_cup,
            "Additional_Notes_cup": Additional_Notes_cup,
            "Movement_Angle_pen": Movement_Angle_pen,
            "Accuracy_pen": Accuracy_pen,
            "Feedback_pen": Feedback_pen,
            "is_adjustment_Required_pen": is_adjustment_Required_pen,
            "Additional_Notes_pen": Additional_Notes_pen,
            "Movement_Angle_key": Movement_Angle_key,
            "Accuracy_key": Accuracy_key,
            "Feedback_key": Feedback_key,
            "is_adjustment_Required_key": is_adjustment_Required_key,
            "Additional_Notes_key": Additional_Notes_key,
            "Movement_Angle_ball": Movement_Angle_ball,
            "Accuracy_ball": Accuracy_ball,
            "Feedback_ball": Feedback_ball,
            "is_adjustment_Required_ball": is_adjustment_Required_ball,
            "Additional_Notes_ball": Additional_Notes_ball,
            "Movement_Angle_bottle": Movement_Angle_bottle,
            "Accuracy_bottle": Accuracy_bottle,
            "Feedback_bottle": Feedback_bottle,
            "is_adjustment_Required_bottle": is_adjustment_Required_bottle,
            "Additional_Notes_bottle": Additional_Notes_bottle,
            "Session_Date": Session_Date,
            "Stimulation_Protocol": Stimulation_Protocol,
            "Target_Movements": Target_Movements,
            "Session_Duration": Session_Duration,
            "Performance_Feedback": Performance_Feedback,
            "Clinician_Notes": Clinician_Notes,
            "Improvements_Observed": Improvements_Observed,
            "Challenges_Noted": Challenges_Noted,
            "Side_effects": Side_effects,
            "Next_Steps_in_Therapy": Next_Steps_in_Therapy,
            "Continue_Current_Protocol": Continue_Current_Protocol,
            "Adjust_Stimulation_Settings": Adjust_Stimulation_Settings,
            "New_Settings": New_Settings,
            "Increase_Session_Frequency": Increase_Session_Frequency,
            "Introduce_New_Exercises": Introduce_New_Exercises,
            "Details": Details,
            "experience_satisfying": experience_satisfying,
            "comfortable_therapy": comfortable_therapy,
            "session_painful": session_painful,
            "hand_improved": hand_improved,
            "Additional_Comments": Additional_Comments,
            "Quality_of_life_Improved": Quality_of_life_Improved,
            "Long_Term_Goals_Achieved": Long_Term_Goals_Achieved,
            "Remaining_Challenges": Remaining_Challenges,
            "Continue_Home_Exercises": Continue_Home_Exercises,
            "Further_Clinical_Sessions_Needed": Further_Clinical_Sessions_Needed,
            "Additional_Therapies_Suggested": Additional_Therapies_Suggested,
        }

    def replace_placeholders(self, doc):
        for paragraph in doc.paragraphs:
            for key, value in self.to_dict().items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in self.to_dict().items():
                        if f"{{{key}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{key}}}", str(value))

    def create_filled_docx(self, output_docx_path):
        doc = Document(TEMPLATE_PATH)
        self.replace_placeholders(doc)
        doc.save(output_docx_path)

    def convert_to_pdf(self, output_docx_path, output_pdf_path):
        convert(output_docx_path, output_pdf_path)

    def upload_to_firebase_storage(self, file_path):
        file_name = os.path.basename(file_path)
        storage.child(file_name).put(file_path)

    def generate_report(self, patient):

        output_pdf_path = (
            f"Add Report to Patient/{self.data["Name"]}_{self.data["Session_Date"]}.pdf"
        )
        self.create_filled_docx(output_docx_path)
        self.convert_to_pdf(output_docx_path, output_pdf_path)
        self.upload_to_firebase_storage(output_pdf_path)

        # Add report to patient and update patient data on Firebase
        patient.add_report(self)
        patient.upload_patient_to_firebase()

        # Delete the DOCX and PDF files after uploading
        if os.path.exists(output_docx_path):
            os.remove(output_docx_path)
        if os.path.exists(output_pdf_path):
            os.remove(output_pdf_path)

    def to_dict(self):
        return self.data


if __name__ == "__main__":
    patient = Patient(
        "Sherif Ahmed",
        30,
        "Male",
        "123-456-7890",
        "987-654-3210",
        "123 Street, City",
        "123ABC",
        "180 cm",
        "75 kg",
        "Right",
        "Left",
        "Flu",
        "Pen",
    )

    report = Report(
        Name="Sherif Ahmed",
        Patient_ID="P002",
        Date_of_Birth="1990-01-01",
        Gender="Male",
        Hand_Dominance="Right",
        Diagnosis="Stroke",
        Referring_Clinician="Dr. Smith",
        Rehabilitation_Start_Date="2023-01-01",
        Muscle_Weakness="Yes",
        Muscle_Paralysis="No",
        Accuracy_ball="",
    )

    report.generate_report(patient)
